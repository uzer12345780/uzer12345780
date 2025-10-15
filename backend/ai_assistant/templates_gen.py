"""
Document Templates Generator
Generates military reports and documents using AI
"""
from flask import Blueprint, request, jsonify
from flask_login import login_required, current_user
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from backend.database.models import db, ActivityLog
from backend.config import TEMP_DIR
import os

templates_bp = Blueprint('templates', __name__)

def generate_report_docx(title, content, author):
    """Generate a Word document report"""
    doc = Document()
    
    # Add title
    doc.add_heading(title, 0)
    
    # Add metadata
    doc.add_paragraph(f"Автор: {author}")
    doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_paragraph("")
    
    # Add content
    for paragraph in content.split('\n\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    
    # Save to temp directory
    filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = TEMP_DIR / filename
    doc.save(str(filepath))
    
    return filepath

@templates_bp.route('/api/templates/report', methods=['POST'])
@login_required
def create_report():
    """Generate a military report document"""
    data = request.get_json()
    
    if not data or not data.get('title') or not data.get('content'):
        return jsonify({'error': 'Title and content are required'}), 400
    
    try:
        title = data.get('title')
        content = data.get('content')
        author = current_user.full_name or current_user.username
        
        filepath = generate_report_docx(title, content, author)
        
        # Log activity
        log = ActivityLog(
            user_id=current_user.id,
            action='report_generated',
            details=f'Generated report: {title}',
            ip_address=request.remote_addr
        )
        db.session.add(log)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Report generated successfully',
            'filename': filepath.name,
            'path': str(filepath)
        }), 200
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@templates_bp.route('/api/templates/list', methods=['GET'])
@login_required
def list_templates():
    """List available document templates"""
    templates = [
        {
            'id': 'report',
            'name': 'Звіт',
            'description': 'Стандартний військовий звіт',
            'fields': ['title', 'content', 'unit', 'date']
        },
        {
            'id': 'request',
            'name': 'Запит',
            'description': 'Запит на матеріали/підтримку',
            'fields': ['title', 'reason', 'items', 'urgency']
        },
        {
            'id': 'order',
            'name': 'Наказ',
            'description': 'Військовий наказ',
            'fields': ['number', 'title', 'content', 'responsible']
        },
        {
            'id': 'briefing',
            'name': 'Брифінг',
            'description': 'Брифінгова записка',
            'fields': ['title', 'situation', 'mission', 'execution']
        }
    ]
    
    return jsonify({'templates': templates}), 200

@templates_bp.route('/api/templates/<template_id>/generate', methods=['POST'])
@login_required
def generate_from_template(template_id):
    """Generate document from a specific template"""
    data = request.get_json()
    
    if not data:
        return jsonify({'error': 'Template data is required'}), 400
    
    try:
        # Template-specific generation logic
        if template_id == 'report':
            content = f"""
ЗВІТ

Підрозділ: {data.get('unit', 'Не вказано')}
Дата: {data.get('date', datetime.now().strftime('%d.%m.%Y'))}

{data.get('content', '')}

Доповідач: {current_user.full_name or current_user.username}
"""
        elif template_id == 'request':
            content = f"""
ЗАПИТ

Підстава: {data.get('reason', '')}

Перелік необхідного:
{data.get('items', '')}

Термін виконання: {data.get('urgency', 'Не вказано')}

Запитувач: {current_user.full_name or current_user.username}
"""
        elif template_id == 'order':
            content = f"""
НАКАЗ №{data.get('number', 'XXX')}

{data.get('title', '')}

{data.get('content', '')}

Відповідальний: {data.get('responsible', '')}

Командир: {current_user.full_name or current_user.username}
"""
        elif template_id == 'briefing':
            content = f"""
БРИФІНГ

Тема: {data.get('title', '')}

1. СИТУАЦІЯ
{data.get('situation', '')}

2. ЗАВДАННЯ
{data.get('mission', '')}

3. ВИКОНАННЯ
{data.get('execution', '')}

Доповідач: {current_user.full_name or current_user.username}
"""
        else:
            return jsonify({'error': 'Unknown template'}), 404
        
        title = data.get('title', f'{template_id.upper()} {datetime.now().strftime("%d.%m.%Y")}')
        filepath = generate_report_docx(title, content, current_user.full_name or current_user.username)
        
        # Log activity
        log = ActivityLog(
            user_id=current_user.id,
            action='template_generated',
            details=f'Generated document from template: {template_id}',
            ip_address=request.remote_addr
        )
        db.session.add(log)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Document generated successfully',
            'filename': filepath.name,
            'path': str(filepath)
        }), 200
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500
